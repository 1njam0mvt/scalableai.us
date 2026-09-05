import logging
from io import BytesIO
from typing import Tuple

logger = logging.getLogger("SCALABLE")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(filename: str, raw_bytes: bytes) -> Tuple[str, str]:
    """Extract plain text from an uploaded file's bytes.

    Returns (text, error). On success `error` is "". On failure `text` is ""
    and `error` explains what went wrong, so the caller can still save the
    file (for the user's record) while marking it as unusable as AI context.
    """

    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return "", f"Unsupported file type '{ext or 'unknown'}'. Supported: PDF, DOCX, TXT, MD."

    try:
        if ext == ".pdf":
            return _extract_pdf(raw_bytes)
        elif ext == ".docx":
            return _extract_docx(raw_bytes)
        else:  # .txt, .md
            return _extract_plain_text(raw_bytes)

    except Exception as e:
        logger.warning("[FILE_PARSER] Extraction failed for %s: %s", filename, e)
        return "", f"Could not read this file: {e}"


def _extract_plain_text(raw_bytes: bytes) -> Tuple[str, str]:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw_bytes.decode(encoding), ""
        except UnicodeDecodeError:
            continue
    return "", "Could not decode this text file (unrecognized encoding)."


def _extract_pdf(raw_bytes: bytes) -> Tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "PDF support isn't installed on the server (missing 'pypdf')."

    reader = PdfReader(BytesIO(raw_bytes))

    if reader.is_encrypted:
        try:
            reader.decrypt("")  # try an empty password - covers PDFs marked encrypted but with no real password
        except Exception:
            return "", "This PDF is password-protected and can't be read."

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception as e:
            logger.warning("[FILE_PARSER] Skipping unreadable PDF page: %s", e)

    text = "\n\n".join(t for t in pages_text if t.strip())

    if not text.strip():
        return "", "No extractable text found (this may be a scanned/image-only PDF)."

    return text, ""


def _extract_docx(raw_bytes: bytes) -> Tuple[str, str]:
    try:
        from docx import Document
    except ImportError:
        return "", "DOCX support isn't installed on the server (missing 'python-docx')."

    doc = Document(BytesIO(raw_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Include table cell text too - tables often carry the actual content
    # (specs, pricing, schedules) that paragraph-only extraction would miss.
    for table in doc.tables:
        for row in table.rows:
            cells_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if cells_text:
                parts.append(cells_text)

    text = "\n".join(parts)

    if not text.strip():
        return "", "No extractable text found in this document."

    return text, ""
